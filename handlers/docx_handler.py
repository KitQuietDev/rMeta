"""
DOCX Metadata Scrubber for rMeta

Reconstructs document text using python-docx to eliminate embedded metadata,
revision history, and author information. Operates in-place with no backups.
"""

import logging
import os
from pathlib import Path

try:
    import docx
except ImportError:
    raise ImportError("python-docx is required. Install with: pip install python-docx")

logger = logging.getLogger(__name__)
__all__ = ["scrub"]

SUPPORTED_EXTENSIONS = {"docx"}

def scrub(file_path: str) -> None:
    """
    Scrubs metadata from a DOCX file by copying paragraph content to a fresh document.

    Args:
        file_path (str): Path to the input DOCX file.

    Raises:
        FileNotFoundError: If the file does not exist.
        PermissionError: If the file cannot be read or written.
        ValueError: If the extension is unsupported.
        RuntimeError: If scrubbing fails or output file is invalid.
    """
    path = Path(file_path)
    ext = path.suffix.lower().lstrip(".")

    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
    if not os.access(file_path, os.R_OK | os.W_OK):
        raise PermissionError(f"Cannot access DOCX file: {file_path}")
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}. Only 'docx' is supported.")

    try:
        doc = docx.Document(file_path)
        new_doc = docx.Document()
        for para in doc.paragraphs:
            new_doc.add_paragraph(para.text)

        temp_path = path.with_suffix(".tmp.docx")
        new_doc.save(temp_path)
        os.replace(temp_path, file_path)
        logger.info(f"📝 DOCX scrubbed: {file_path}")

    except Exception as e:
        logger.error(f"❌ Error scrubbing DOCX metadata: {e}")
        raise RuntimeError(f"Failed to scrub DOCX: {e}")

    if not path.exists() or path.stat().st_size == 0:
        raise RuntimeError(f"Scrubbed DOCX file missing or empty: {file_path}")
