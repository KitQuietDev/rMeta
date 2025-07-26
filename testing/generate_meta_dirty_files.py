import os
import csv
import random
import logging
from pathlib import Path
from datetime import datetime
from PIL import Image, ImageDraw, PngImagePlugin
try:
    import pillow_heif  # Optional: enables HEIC decoding via Pillow
except ImportError:
    pillow_heif = None

from docx import Document
from docx.shared import Inches
from PyPDF2 import PdfWriter
from openpyxl import Workbook
import piexif  # Used for injecting EXIF metadata into JPGs

# === OUTPUT SETUP ===

OUTPUT_DIR = Path("meta_dirty_bundle")  # Directory for all generated corrupt files
CORRUPTION_LOG = OUTPUT_DIR / "corruption_log.txt"  # Log file capturing corruption details
HEIC_FILE = Path("clean.heic")  # Clean HEIC source used for decoding test

# === HELPER SETUP ===

def setup_logging():
    # Ensure log directory exists before logging starts
    CORRUPTION_LOG.parent.mkdir(parents=True, exist_ok=True)

    # Configure logging to write timestamped events
    logging.basicConfig(filename=CORRUPTION_LOG, level=logging.INFO,
                        format="%(asctime)s: %(message)s")
    logging.info("Starting metadata corruption gauntlet")

def ensure_dir():
    # Create output directory if it doesn't already exist
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# === FILE CORRUPTION FUNCTIONS ===

def corrupt_txt():
    # Simple TXT file with embedded fake PII and manipulated timestamp
    path = OUTPUT_DIR / "dirty.txt"
    with open(path, "w") as f:
        f.write("Name: John Doe\nEmail: john@example.com\nSSN: 123-45-6789\n")

    # Spoof file timestamp to Jan 1, 2021 for audit simulation
    os.utime(path, (1609459200, 1609459200))
    logging.info("Corrupted TXT with PII and spoofed timestamp")

def corrupt_csv():
    # CSV populated with synthetic PII rows
    path = OUTPUT_DIR / "dirty.csv"
    with open(path, "w", newline='') as f:
        writer = csv.writer(f)
        writer.writerow(["Name", "Email", "SSN"])
        for i in range(3):
            writer.writerow([
                f"Jane{i}", f"jane{i}@example.com",
                f"{random.randint(100,999)}-{random.randint(10,99)}-{random.randint(1000,9999)}"
            ])
    logging.info("Corrupted CSV with synthetic PII rows")

def corrupt_png():
    # PNG with visible overlay and injected metadata
    img = Image.new("RGB", (100, 100), color="red")
    draw = ImageDraw.Draw(img)
    draw.text((10, 40), "Metadata!", fill="white")

    # Inject fake GPS and junk comment into PNG header
    meta = PngImagePlugin.PngInfo()
    meta.add_text("GPS", "N 48.8588443, E 2.2943506")
    meta.add_text("Comment", "Steganographic junk embedded here")
    
    img_path = OUTPUT_DIR / "dirty.png"
    img.save(img_path, "PNG", pnginfo=meta)
    logging.info("Created PNG with spoofed GPS + comment metadata")

def corrupt_docx():
    # DOCX with hidden SSN and spoofed author metadata
    path = OUTPUT_DIR / "dirty.docx"
    doc = Document()
    doc.add_paragraph("This document contains hidden metadata.")

    # Insert fake SSN as a separate paragraph
    p = doc.add_paragraph()
    run = p.add_run("SSN: 000-00-0000")
    run.font.color.rgb = None  # No visibility manipulation here; just neutral color

    # Overwrite author metadata with fake ID
    doc.core_properties.author = "EvilDocBot"
    doc.save(path)
    logging.info("Corrupted DOCX with hidden SSN and spoofed author metadata")

def corrupt_pdf():
    # PDF with blank page and injected metadata keys (including junk)
    path = OUTPUT_DIR / "dirty.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)

    # Set standard and non-standard metadata
    writer.add_metadata({
        "/Author": "GhostWriter",
        "/Title": "Corrupted Metadata",
        "/Subject": "Hidden Data Test",
        "/JunkKey": "\x00\x01\x02\x03"  # Invalid characters for stress testing
    })

    with open(path, "wb") as f:
        writer.write(f)
    logging.info("Corrupted PDF with metadata payload including junk keys")

def corrupt_xlsx():
    # XLSX with custom sheet name and embedded email addresses
    path = OUTPUT_DIR / "dirty.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "MetadataSoup"
    ws["A1"] = "Name"
    ws["B1"] = "Email"
    ws["A2"] = "Alice"
    ws["B2"] = "alice@example.com"
    wb.save(path)
    logging.info("Corrupted XLSX with custom sheet name and email field")

def corrupt_jpg():
    # JPG with visible text overlay and spoofed EXIF metadata
    path = OUTPUT_DIR / "dirty.jpg"
    img = Image.new("RGB", (120, 120), color="blue")
    draw = ImageDraw.Draw(img)
    draw.text((10, 50), "ExifBomb!", fill="white")

    # Construct minimal EXIF payload
    zeroth_ifd = {
        piexif.ImageIFD.Make: u"Canon",
        piexif.ImageIFD.Model: u"Corruptomatic 3000",
        piexif.ImageIFD.Software: u"piexif"
    }
    exif_dict = {"0th": zeroth_ifd}
    exif_bytes = piexif.dump(exif_dict)

    img.save(path, exif=exif_bytes)
    logging.info("Corrupted JPG with spoofed EXIF camera info")

def corrupt_heic():
    # Optional: convert HEIC to JPEG (if Pillow can’t encode HEIC directly)
    if not pillow_heif:
        logging.warning("HEIC handler not installed — skipping")
        return

    if not HEIC_FILE.exists():
        logging.warning("Missing clean.heic file — skipping HEIC corruption")
        return

    img = pillow_heif.read_heif(HEIC_FILE)[0].to_pillow()
    path = OUTPUT_DIR / "dirty.heic.jpg"
    img.save(path, format="JPEG")  # HEIC encoding unsupported, fallback to JPEG
    logging.info("Converted HEIC to JPEG for output")

# === EXECUTION DRIVER ===

def run_gauntlet():
    ensure_dir()       # Create output folder
    setup_logging()    # Start logging after folder exists
    corrupt_txt()      # Begin file creation (one per format)
    corrupt_csv()
    corrupt_png()
    corrupt_docx()
    corrupt_pdf()
    corrupt_xlsx()
    corrupt_jpg()
    corrupt_heic()     # Optional, dependent on install and file presence

# === SCRIPT ENTRY POINT ===

if __name__ == "__main__":
    run_gauntlet()  # Launch the gauntlet when script is executed directly
